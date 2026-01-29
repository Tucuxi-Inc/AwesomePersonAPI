"""Resume Parser service for extracting text from PDF and DOCX files."""

import io
import logging
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


class ResumeParserError(Exception):
    """Error during resume parsing."""
    pass


class ResumeParser:
    """Service for extracting text content from resume files."""

    SUPPORTED_TYPES = {"pdf", "docx", "doc", "txt"}

    def parse(self, file_content: bytes, file_type: str) -> str:
        """
        Extract text from a resume file.

        Args:
            file_content: Raw file bytes
            file_type: File extension (pdf, docx, etc.)

        Returns:
            Extracted text content

        Raises:
            ResumeParserError: If parsing fails
        """
        file_type = file_type.lower().lstrip(".")

        if file_type not in self.SUPPORTED_TYPES:
            raise ResumeParserError(f"Unsupported file type: {file_type}")

        try:
            if file_type == "pdf":
                return self._parse_pdf(file_content)
            elif file_type in ("docx", "doc"):
                return self._parse_docx(file_content)
            elif file_type == "txt":
                return self._parse_txt(file_content)
            else:
                raise ResumeParserError(f"Unsupported file type: {file_type}")
        except ResumeParserError:
            raise
        except Exception as e:
            logger.error(f"Error parsing resume: {e}")
            raise ResumeParserError(f"Failed to parse resume: {str(e)}")

    def _parse_pdf(self, content: bytes) -> str:
        """Extract text from PDF using pypdf."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ResumeParserError("pypdf package not installed")

        try:
            reader = PdfReader(io.BytesIO(content))
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise ResumeParserError(
                    "PDF appears to be scanned or image-based. "
                    "Please upload a text-based PDF or DOCX file."
                )

            return self._clean_text(full_text)

        except ResumeParserError:
            raise
        except Exception as e:
            raise ResumeParserError(f"Failed to parse PDF: {str(e)}")

    def _parse_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ResumeParserError("python-docx package not installed")

        try:
            doc = Document(io.BytesIO(content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n".join(text_parts)

            if not full_text.strip():
                raise ResumeParserError("DOCX file appears to be empty.")

            return self._clean_text(full_text)

        except ResumeParserError:
            raise
        except Exception as e:
            raise ResumeParserError(f"Failed to parse DOCX: {str(e)}")

    def _parse_txt(self, content: bytes) -> str:
        """Extract text from plain text file."""
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin-1")

            if not text.strip():
                raise ResumeParserError("Text file appears to be empty.")

            return self._clean_text(text)

        except ResumeParserError:
            raise
        except Exception as e:
            raise ResumeParserError(f"Failed to parse text file: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Clean up extracted text."""
        # Remove excessive whitespace
        lines = []
        for line in text.split("\n"):
            # Remove leading/trailing whitespace
            line = line.strip()
            # Collapse multiple spaces
            line = " ".join(line.split())
            if line:
                lines.append(line)

        # Remove duplicate blank lines
        cleaned_lines = []
        prev_blank = False
        for line in lines:
            if not line:
                if not prev_blank:
                    cleaned_lines.append("")
                prev_blank = True
            else:
                cleaned_lines.append(line)
                prev_blank = False

        return "\n".join(cleaned_lines)

    def detect_file_type(self, filename: str) -> str:
        """Detect file type from filename."""
        ext = Path(filename).suffix.lower().lstrip(".")
        return ext if ext in self.SUPPORTED_TYPES else "unknown"

    def validate_file(
        self, content: bytes, filename: str, max_size_mb: float = 10
    ) -> Tuple[bool, Optional[str]]:
        """
        Validate a resume file.

        Args:
            content: File content
            filename: Original filename
            max_size_mb: Maximum file size in MB

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file size
        size_mb = len(content) / (1024 * 1024)
        if size_mb > max_size_mb:
            return False, f"File too large. Maximum size is {max_size_mb}MB."

        # Check file type
        file_type = self.detect_file_type(filename)
        if file_type not in self.SUPPORTED_TYPES:
            return False, f"Unsupported file type. Supported: {', '.join(self.SUPPORTED_TYPES)}"

        # Check content is not empty
        if len(content) < 100:
            return False, "File appears to be empty or too small."

        return True, None


# Singleton instance
resume_parser = ResumeParser()


def get_resume_parser() -> ResumeParser:
    """Get the resume parser instance."""
    return resume_parser
